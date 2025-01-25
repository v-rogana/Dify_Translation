from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document
import os

def markdown_to_docx(input_md_file, output_docx_file):
    try:
        # Abrir o arquivo Markdown
        with open(input_md_file, 'r', encoding='utf-8') as file:
            md_content = file.read()
        
        # Converter Markdown para HTML
        html_content = markdown(md_content)
        
        # Parsear o HTML com BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Criar um documento Word
        doc = Document()
        
        # Iterar pelos elementos do HTML e adicioná-los ao documento
        for element in soup.descendants:
            if element.name == 'h1':  # Título nível 1
                doc.add_heading(element.text, level=1)
            elif element.name == 'h2':  # Título nível 2
                doc.add_heading(element.text, level=2)
            elif element.name == 'h3':  # Título nível 3
                doc.add_heading(element.text, level=3)
            elif element.name == 'p':  # Parágrafos
                doc.add_paragraph(element.text)
            elif element.name == 'code' and element.parent.name == 'pre':  # Blocos de código
                pre = doc.add_paragraph()
                pre.add_run(element.text).font.name = 'Courier New'
            elif element.name == 'li':  # Listas
                doc.add_paragraph(f'• {element.text}', style='List Bullet')
        
        # Salvar o documento Word
        doc.save(output_docx_file)
        print(f"Arquivo convertido para {output_docx_file}")
    except Exception as e:
        print(f"Erro ao converter o arquivo: {e}")

if __name__ == "__main__":
    input_md_file = os.path.join(os.path.dirname(__file__), "..", "data", "output", "translated", "unified_output.md")
    output_docx_file = os.path.join(os.path.dirname(__file__), "..", "data", "output", "translated", "unified_output.docx")

    markdown_to_docx(input_md_file, output_docx_file)
